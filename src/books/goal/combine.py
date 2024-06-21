folder_path = '/Users/farshid/code/Projects/pirahansiah.github.io/src/books/goal'
output_file = '/Users/farshid/code/Projects/pirahansiah.github.io/src/books/goal/combined_document.docx'

import os
from docx import Document
from docx.shared import Inches

def combine_docx_files(folder_path, output_file):
    combined_document = Document()

    def remove_footer(doc):
        for section in doc.sections:
            sectPr = section._sectPr
            for element in sectPr:
                if element.tag.endswith(('footerReference', 'pgNumType')):
                    sectPr.remove(element)

    def add_image(doc, image_path):
        if os.path.exists(image_path):
            # Create a new paragraph and add the image
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(6.0))  # Adjust width as needed

    for subdir, _, files in os.walk(folder_path):
        subdir_name = os.path.basename(subdir).lower()
        image_path = os.path.join(subdir, f"{subdir_name}1.png")
        
        for file in files:
            if file.endswith(".docx"):
                file_path = os.path.join(subdir, file)
                doc = Document(file_path)
                remove_footer(doc)

                # Create a temporary document to hold the image and then the original document's content
                temp_doc = Document()
                add_image(temp_doc, image_path)
                for element in doc.element.body:
                    temp_doc.element.body.append(element)
                
                # Append the temporary document's content to the combined document
                for element in temp_doc.element.body:
                    combined_document.element.body.append(element)

    combined_document.save(output_file)

combine_docx_files(folder_path, output_file)
